from flask import Flask, request, render_template, jsonify, send_file, abort
from docx import Document
from typing import Dict, Optional, List
import os
import logging
from pathlib import Path
import traceback

app = Flask(__name__)

# Configure detailed logging
logging.basicConfig(
    level=logging.DEBUG,
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s'
)
logger = logging.getLogger(__name__)


class ContractState:
    def __init__(self):
        self.step: int = 1
        self.responses: Dict[str, str] = {}
        self.user_type: str = ""
        self.contract_type: str = ""

    def reset(self):
        self.__init__()


# Configuration
VALID_ANSWERS = {
    "user_types": ["임대인", "임차인"],
    "locations": ["서울특별시", "부산광역시", "대구광역시", "인천광역시", "대전광역시"],
    "contract_types": ["월세", "전세"],
    "payment_types": ["보증금+월세", "보증금 없는 월세", "전액 보증금", "반전세"]
}

conversation_state = ContractState()


class ContractGenerator:
    def __init__(self):
        current_file = os.path.abspath(__file__)
        self.base_dir = os.path.dirname(current_file)
        self.output_dir = os.path.join(self.base_dir, "../../generated")
        self.template_dir = os.path.join(self.base_dir, "..")

        logger.info(f"Base directory: {self.base_dir}")
        logger.info(f"Output directory: {self.output_dir}")
        logger.info(f"Template directory: {self.template_dir}")

        self.setup_directories()

    def setup_directories(self):
        """Ensure necessary directories exist"""
        try:
            if not os.path.exists(self.output_dir):
                os.makedirs(self.output_dir)
                logger.info(f"Created output directory: {self.output_dir}")

            if not os.path.exists(self.template_dir):
                os.makedirs(self.template_dir)
                logger.info(f"Created template directory: {self.template_dir}")

            template_path = os.path.join(self.template_dir, "간이임대차계약서(placeholder).docx")
            if not os.path.exists(template_path):
                logger.error(f"Template file not found at: {template_path}")
                raise FileNotFoundError(f"Contract template file not found at {template_path}")
            else:
                logger.info("Template file found successfully")

        except Exception as e:
            logger.error(f"Error in setup_directories: {str(e)}")
            logger.error(traceback.format_exc())
            raise

    def replace_text_with_bold(self, paragraph, placeholder: str, new_text: str):
        """Replace placeholder text with bold formatted text"""
        if placeholder in paragraph.text:
            # 기존 텍스트를 분할
            parts = paragraph.text.split(placeholder)

            # 단락의 모든 runs를 제거
            for run in paragraph.runs:
                run._element.getparent().remove(run._element)

            # 새로운 텍스트로 재구성
            for i, part in enumerate(parts):
                # 분할된 일반 텍스트 추가
                if part:
                    run = paragraph.add_run(part)

                # 마지막 부분이 아니면 placeholder를 대체할 새 텍스트 추가
                if i < len(parts) - 1:
                    bold_run = paragraph.add_run(new_text)
                    bold_run.bold = True  # 굵은 글씨 적용

    def generate(self, responses: Dict[str, str], user_type: str, contract_type: str) -> str:
        """Generate contract document based on user responses"""
        try:
            template_path = os.path.join(self.template_dir, "간이임대차계약서(placeholder).docx")
            output_filename = "완성된_임대차계약서.docx"
            output_path = os.path.join(self.output_dir, output_filename)

            logger.info(f"Loading template from: {template_path}")
            logger.info(f"Will save output to: {output_path}")
            logger.info(f"Responses received: {responses}")

            # Load the document
            doc = Document(template_path)
            logger.info("Template document loaded successfully")

            # Prepare placeholders with their replacement values
            placeholders = {
                "{p1}": responses.get("Q3", ""),  # 부동산 소재지
                "{p21}" if user_type == "임대인" else "{p22}": responses.get("Q1", ""),  # 성명
                "{p3}": responses.get("Q4", ""),  # 보증금 금액
                "{p4}": responses.get("Q5", "")  # 월세 금액
            }

            logger.info(f"Replacing placeholders with bold text: {placeholders}")

            # Replace text in paragraphs with bold formatting
            for paragraph in doc.paragraphs:
                original_text = paragraph.text
                for placeholder, new_text in placeholders.items():
                    if placeholder in paragraph.text:
                        self.replace_text_with_bold(paragraph, placeholder, new_text)
                        logger.info(f"Replaced '{placeholder}' with bold '{new_text}' in text")

            # Save the document
            logger.info(f"Saving document to: {output_path}")
            doc.save(output_path)
            logger.info("Document saved successfully")

            return output_filename

        except Exception as e:
            logger.error(f"Contract generation failed: {str(e)}")
            logger.error(traceback.format_exc())
            raise


class InputValidator:
    @staticmethod
    def validate_name(name: str) -> bool:
        """Validate user name input"""
        return bool(name and name.isalpha())

    @staticmethod
    def validate_amount(amount: str) -> bool:
        """Validate monetary amount input"""
        return bool(amount and amount.isdigit())

    @staticmethod
    def validate_option(value: str, valid_options: List[str]) -> bool:
        """Validate if input is in allowed options"""
        return value in valid_options


contract_generator = ContractGenerator()


@app.route("/")
def home():
    return render_template("index.html")


@app.route("/chat", methods=["POST"])
def chat():
    try:
        user_message = request.json.get("message", "").strip()
        if not user_message:
            return jsonify({"error": "메시지를 입력해주세요."})

        step = conversation_state.step
        response = ""
        options = None

        if step == 1:
            response = "1. 임대인/임차인 여부를 선택해 주세요 (임대인 / 임차인)"
            options = VALID_ANSWERS["user_types"]
            conversation_state.step += 1

        elif step == 2:
            if not InputValidator.validate_option(user_message, VALID_ANSWERS["user_types"]):
                return jsonify({"error": "임대인 또는 임차인 중에서 선택해주세요."})
            conversation_state.user_type = user_message
            response = "2. 임대인/임차인의 성명을 입력해주세요."
            conversation_state.step += 1

        elif step == 3:
            if not InputValidator.validate_name(user_message):
                return jsonify({"error": "이름을 올바로 입력해주세요."})
            conversation_state.responses["Q1"] = user_message
            response = "3. 부동산 소재지는 어떻게 되나요?"
            options = VALID_ANSWERS["locations"]
            conversation_state.step += 1

        elif step == 4:
            if not InputValidator.validate_option(user_message, VALID_ANSWERS["locations"]):
                return jsonify({"error": f"잘못된 입력입니다. 다음 중에서 선택해 주세요: {', '.join(VALID_ANSWERS['locations'])}"})
            conversation_state.responses["Q3"] = user_message
            response = "4. 월세 또는 전세 여부를 선택해 주세요"
            options = VALID_ANSWERS["contract_types"]
            conversation_state.step += 1

        elif step == 5:
            if not InputValidator.validate_option(user_message, VALID_ANSWERS["contract_types"]):
                return jsonify({"error": "월세 또는 전세 중에서 선택해주세요."})
            conversation_state.contract_type = user_message
            response = "5. 보증금 금액을 입력해주세요 (숫자만 입력)"
            conversation_state.step += 1

        elif step == 6:
            if not InputValidator.validate_amount(user_message):
                return jsonify({"error": "올바른 금액을 입력해주세요 (숫자만 입력)."})
            conversation_state.responses["Q4"] = user_message
            response = "6. 월세 금액을 입력해주세요 (숫자만 입력)"
            conversation_state.step += 1

        elif step == 7:
            if not InputValidator.validate_amount(user_message):
                return jsonify({"error": "올바른 금액을 입력해주세요 (숫자만 입력)."})
            conversation_state.responses["Q5"] = user_message

            try:
                output_filename = contract_generator.generate(
                    conversation_state.responses,
                    conversation_state.user_type,
                    conversation_state.contract_type
                )
                logger.info(f"Contract generated successfully: {output_filename}")

                conversation_state.reset()
                return jsonify({
                    "response": "감사합니다! 완성된 계약서 파일을 확인해 주세요.",
                    "file_url": f"/download/{output_filename}"
                })
            except Exception as e:
                logger.error(f"Contract generation failed: {str(e)}")
                logger.error(traceback.format_exc())
                return jsonify({"error": "계약서 생성 중 오류가 발생했습니다. 다시 시도해주세요."})

    except Exception as e:
        logger.error(f"Unexpected error: {str(e)}")
        logger.error(traceback.format_exc())
        conversation_state.reset()
        return jsonify({"error": "처리 중 오류가 발생했습니다. 다시 시도해주세요."})

    return jsonify({"response": response, "options": options})


@app.route("/download/<filename>")
def download_file(filename):
    try:
        file_path = os.path.join(contract_generator.output_dir, filename)
        logger.info(f"Attempting to download file: {file_path}")

        if not os.path.exists(file_path):
            logger.error(f"File not found: {file_path}")
            abort(404)

        return send_file(file_path, as_attachment=True)
    except Exception as e:
        logger.error(f"Download failed: {str(e)}")
        logger.error(traceback.format_exc())
        abort(500)


if __name__ == "__main__":
    app.run(debug=True)