from flask import Flask, request, render_template, jsonify, send_file, abort
from docx import Document
from typing import Dict, Optional, List
import os
import logging
from pathlib import Path
import traceback
from openai import OpenAI
import json

os.environ["OPENAI_API_KEY"] = "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"

app = Flask(__name__)
logger = logging.getLogger(__name__)

conversation_state = None

SPECIAL_TERMS_FILE = "data/special_terms.json"
USER_RESPONSES_FILE = "data/user_responses.json"

def initialize_conversation_state():
    """대화 상태 초기화 함수"""
    global conversation_state
    if conversation_state is None:
        conversation_state = ContractState()
    return conversation_state

def save_special_term(term, recommendation):
    os.makedirs(os.path.dirname(SPECIAL_TERMS_FILE), exist_ok=True)
    if os.path.exists(SPECIAL_TERMS_FILE):
        with open(SPECIAL_TERMS_FILE, "r", encoding="utf-8") as file:
            saved_terms = json.load(file)
    else:
        saved_terms = {}
    saved_terms[term] = recommendation
    with open(SPECIAL_TERMS_FILE, "w", encoding="utf-8") as file:
        json.dump(saved_terms, file, ensure_ascii=False, indent=4)


def save_user_responses():
    """사용자 응답 저장 함수"""
    try:
        global conversation_state
        if conversation_state is None:
            logger.error("Conversation state is None")
            return

        folder_path = os.path.dirname(USER_RESPONSES_FILE)
        os.makedirs(folder_path, exist_ok=True)

        data = {
            "responses": conversation_state.responses,
            "user_type": conversation_state.user_type,
            "contract_type": conversation_state.contract_type
        }

        logger.info(f"Saving user responses: {data}")

        with open(USER_RESPONSES_FILE, "w", encoding="utf-8") as file:
            json.dump(data, file, ensure_ascii=False, indent=4)

        logger.info(f"✅ 사용자 입력값이 {USER_RESPONSES_FILE}에 저장되었습니다.")

    except Exception as e:
        logger.error(f"Error saving user responses: {str(e)}")
        raise

def load_user_responses():
    """사용자 응답 데이터를 로드하는 함수"""
    try:
        if os.path.exists(USER_RESPONSES_FILE):
            with open(USER_RESPONSES_FILE, "r", encoding="utf-8") as file:
                data = json.load(file)
                logger.info(f"Successfully loaded user responses: {data}")
                return data
        else:
            logger.warning("User responses file not found")
            return {"responses": {}, "user_type": "", "contract_type": ""}
    except Exception as e:
        logger.error(f"Error loading user responses: {str(e)}")
        return {"responses": {}, "user_type": "", "contract_type": ""}


class GPTConnector:
    @staticmethod
    def get_special_terms_recommendations(term_type: str) -> dict:
        try:
            client = OpenAI(api_key=os.getenv("OPENAI_API_KEY"))
            prompt = f"'{term_type}' 관련하여 임대차계약서에 넣으면 좋을 만한 특약사항을 추천해줘."
            response = client.chat.completions.create(
                model="gpt-4o-mini",
                messages=[
                    {"role": "system", "content": "You are a helpful assistant."},
                    {"role": "user", "content": prompt}
                ],
                max_tokens=300,
                temperature=0.7
            )
            recommendations_text = response.choices[0].message.content.strip()
            split_parts = split_text_into_parts(recommendations_text)
            return split_parts
        except Exception as e:
            logger.error(f"GPT recommendation failed for term type '{term_type}': {str(e)}")
            return {"head": "추천 항목을 불러오지 못했습니다. 다시 시도해주세요.", "item1": "", "item2": ""}


def split_text_into_parts(full_text: str) -> dict:
    head_index = full_text.find("1.")
    item2_index = full_text.find("2.")
    item3_index = full_text.find("3.")
    head = full_text[:head_index].strip() if head_index != -1 else full_text
    item1 = full_text[head_index:item2_index].strip() if head_index != -1 and item2_index != -1 else ""
    if item2_index != -1:
        if item3_index != -1:
            item2 = full_text[item2_index:item3_index].strip()
        else:
            item2 = full_text[item2_index:].strip()
    else:
        item2 = ""
    return {"head": head, "item1": item1, "item2": item2}


class ContractState:
    def __init__(self):
        self.step: int = 1
        saved_data = load_user_responses()
        self.responses: Dict[str, str] = saved_data.get("responses", {})
        self.user_type: str = saved_data.get("user_type", "")
        self.contract_type: str = saved_data.get("contract_type", "")

    def reset(self):
        self.__init__()
        if os.path.exists(USER_RESPONSES_FILE):
            os.remove(USER_RESPONSES_FILE)

# Configuration
VALID_ANSWERS = {
    "user_types": ["임대인", "임차인"],
    "locations": ["서울특별시", "부산광역시", "대구광역시", "인천광역시", "대전광역시"],
    "contract_types": ["월세", "전세"],
    "payment_types": ["보증금+월세", "보증금 없는 월세", "전액 보증금", "반전세"]
}

conversation_state = ContractState()

class ContractGenerator:
    def __init__(self):
        current_file = os.path.abspath(__file__)
        self.base_dir = os.path.dirname(current_file)
        self.output_dir = os.path.join(self.base_dir, "generated")
        self.template_dir = os.path.join(self.base_dir, "templates")

        logger.info(f"Base directory: {self.base_dir}")
        logger.info(f"Output directory: {self.output_dir}")
        logger.info(f"Template directory: {self.template_dir}")

        self.setup_directories()

    def setup_directories(self):
        try:
            if not os.path.exists(self.output_dir):
                os.makedirs(self.output_dir)
                logger.info(f"Created output directory: {self.output_dir}")

            if not os.path.exists(self.template_dir):
                os.makedirs(self.template_dir)
                logger.info(f"Created template directory: {self.template_dir}")

            template_path = os.path.join(self.template_dir, "간이임대차계약서(placeholder).docx")
            if not os.path.exists(template_path):
                logger.error(f"Template file not found at: {template_path}")
                raise FileNotFoundError(f"Contract template file not found at {template_path}")
            else:
                logger.info("Template file found successfully")

        except Exception as e:
            logger.error(f"Error in setup_directories: {str(e)}")
            logger.error(traceback.format_exc())
            raise

    def preserve_text_format(self, paragraph, placeholder: str, new_text: str):
        if placeholder not in paragraph.text:
            return

        formatted_runs = []
        placeholder_run = None
        for run in paragraph.runs:
            if placeholder in run.text:
                placeholder_run = run
                format_info = {
                    'bold': run.bold,
                    'italic': run.italic,
                    'underline': run.underline,
                    'font': run.font.name if run.font.name else None,
                    'size': run.font.size if run.font.size else None,
                    'highlight_color': run._element.rPr.highlight_val if hasattr(run._element.rPr, 'highlight_val') else None,
                    'color': run._element.rPr.color_val if hasattr(run._element.rPr, 'color_val') else None,
                }
                formatted_runs.append((run.text, format_info))

        if placeholder_run and formatted_runs:
            start_text = placeholder_run.text.split(placeholder)[0]
            end_text = placeholder_run.text.split(placeholder)[1]

            placeholder_run._element.getparent().remove(placeholder_run._element)

            if start_text:
                new_run = paragraph.add_run(start_text)
                self._apply_format(new_run, formatted_runs[0][1])

            new_run = paragraph.add_run(new_text)
            self._apply_format(new_run, formatted_runs[0][1])

            if end_text:
                new_run = paragraph.add_run(end_text)
                self._apply_format(new_run, formatted_runs[0][1])

    def _apply_format(self, run, format_info):
        run.bold = format_info['bold']
        run.italic = format_info['italic']
        run.underline = format_info['underline']
        if format_info['font']:
            run.font.name = format_info['font']
        if format_info['size']:
            run.font.size = format_info['size']

    def generate(self, responses: Dict[str, str], user_type: str, contract_type: str, selected_terms: List[str]) -> str:
        try:
            template_path = os.path.join(self.template_dir, "간이임대차계약서(placeholder).docx")
            output_filename = "완성된_임대차계약서.docx"
            output_path = os.path.join(self.output_dir, output_filename)

            doc = Document(template_path)

            # 🔥 JSON 데이터 기반으로 플레이스홀더 설정
            placeholders = {
                "p1": responses.get("location", ""),  # 소재지
                "p21" if user_type == "임대인" else "p22": responses.get("user_name", ""),  # 임대인/임차인 이름
                "p3": responses.get("deposit", ""),  # 보증금
                "p4": responses.get("monthly_rent", ""),  # 월세
            }

            # 🔥 특약사항 반영
            for i, term in enumerate(selected_terms[:4]):
                placeholders[f"result{i + 1}"] = term

            # 🔥 플레이스홀더를 계약서에 반영
            for paragraph in doc.paragraphs:
                for placeholder, new_text in placeholders.items():
                    self.preserve_text_format(paragraph, f"{{{placeholder}}}", new_text)

            doc.save(output_path)
            logger.info(f"✅ 계약서가 {output_path}에 저장되었습니다.")
            return output_filename

        except Exception as e:
            logger.error(f"❌ 계약서 생성 실패: {str(e)}")
            raise


class InputValidator:
    @staticmethod
    def validate_name(name: str) -> bool:
        return bool(name and name.isalpha())

    @staticmethod
    def validate_amount(amount: str) -> bool:
        return bool(amount and amount.isdigit())

    @staticmethod
    def validate_option(value: str, valid_options: List[str]) -> bool:
        return value in valid_options


contract_generator = ContractGenerator()


@app.route("/special_terms", methods=["POST"])
def special_terms():
    try:
        # 현재 상태 유지
        global conversation_state
        initialize_conversation_state()

        selected_terms = request.json.get("selected_terms", [])
        if not selected_terms:
            return jsonify({"error": "특약 사항 유형을 선택해주세요."})

        recommendations = {}
        for term in selected_terms:
            recommendation = GPTConnector.get_special_terms_recommendations(term)
            recommendations[term] = recommendation
            save_special_term(term, recommendation)

        # 현재 상태를 파일에 저장
        save_user_responses()

        return jsonify({
            "message": "추천된 특약사항들을 확인해 보세요. 체크를 표시하는 특약 사항들은 계약서에 반영됩니다!",
            "recommendations": recommendations
        })
    except Exception as e:
        logger.error(f"Error in special_terms: {str(e)}")
        return jsonify({"error": "특약 사항 추천을 불러오는 중 문제가 발생했습니다. 다시 시도해주세요."})


@app.route("/finalize_terms", methods=["POST"])
def finalize_terms():
    try:
        # 저장된 상태 로드
        saved_data = load_user_responses()
        logger.info(f"Loaded user responses: {saved_data}")

        if not saved_data or not saved_data.get("responses"):
            return jsonify({"error": "사용자 입력 데이터를 찾을 수 없습니다. 처음부터 다시 시작해주세요."})

        selected_terms = request.json.get("selected_terms", [])
        if not selected_terms:
            return jsonify({"error": "특약사항을 선택해주세요."})

        # generate 함수 호출 시 저장된 데이터 전달
        output_filename = contract_generator.generate(
            responses=saved_data["responses"],
            user_type=saved_data["user_type"],
            contract_type=saved_data["contract_type"],
            selected_terms=selected_terms
        )

        return jsonify({
            "message": "선택한 특약사항이 계약서에 반영되었습니다.",
            "file": output_filename
        })

    except Exception as e:
        logger.error(f"Finalize terms failed: {str(e)}")
        logger.error(traceback.format_exc())
        return jsonify({"error": "특약사항 반영 중 오류가 발생했습니다. 다시 시도해주세요."})



@app.route("/")
def home():
    # 앱 시작시 상태 초기화
    global conversation_state
    conversation_state = ContractState()
    return render_template("index.html")


@app.route("/chat", methods=["POST"])
def chat():
    try:
        user_message = request.json.get("message", "").strip()
        if not user_message:
            return jsonify({"error": "메시지를 입력해주세요."})

        step = conversation_state.step
        response = ""
        options = None

        if step == 1:
            response = "1. 임대인/임차인 여부를 선택해 주세요 (임대인 / 임차인)"
            options = VALID_ANSWERS["user_types"]
            conversation_state.step += 1

        elif step == 2:
            if not InputValidator.validate_option(user_message, VALID_ANSWERS["user_types"]):
                return jsonify({"error": "임대인 또는 임차인 중에서 선택해주세요."})
            conversation_state.user_type = user_message
            save_user_responses()  # 입력값 저장
            response = "2. 임대인/임차인의 성명을 입력해주세요."
            conversation_state.step += 1

        elif step == 3:
            if not InputValidator.validate_name(user_message):
                return jsonify({"error": "이름을 올바로 입력해주세요."})
            conversation_state.responses["user_name"] = user_message
            save_user_responses()  # 입력값 저장
            response = "3. 부동산 소재지는 어떻게 되나요?"
            options = VALID_ANSWERS["locations"]
            conversation_state.step += 1

        elif step == 4:
            if not InputValidator.validate_option(user_message, VALID_ANSWERS["locations"]):
                return jsonify({"error": f"잘못된 입력입니다. 다음 중에서 선택해 주세요: {', '.join(VALID_ANSWERS['locations'])}"})
            conversation_state.responses["location"] = user_message
            save_user_responses()  # 입력값 저장
            response = "4. 월세 또는 전세 여부를 선택해 주세요"
            options = VALID_ANSWERS["contract_types"]
            conversation_state.step += 1

        elif step == 5:
            if not InputValidator.validate_option(user_message, VALID_ANSWERS["contract_types"]):
                return jsonify({"error": "월세 또는 전세 중에서 선택해주세요."})
            conversation_state.contract_type = user_message
            save_user_responses()  # 입력값 저장
            response = "5. 보증금 금액을 입력해주세요 (숫자만 입력)"
            conversation_state.step += 1

        elif step == 6:
            if not InputValidator.validate_amount(user_message):
                return jsonify({"error": "올바른 금액을 입력해주세요 (숫자만 입력)."})
            conversation_state.responses["deposit"] = user_message
            save_user_responses()  # 입력값 저장
            response = "6. 월세 금액을 입력해주세요 (숫자만 입력)"
            conversation_state.step += 1

        elif step == 7:
            if not InputValidator.validate_amount(user_message):
                return jsonify({"error": "올바른 금액을 입력해주세요 (숫자만 입력)."})
            conversation_state.responses["monthly_rent"] = user_message
            save_user_responses()  # 입력값 저장

            try:
                output_filename = contract_generator.generate(
                    conversation_state.responses,
                    conversation_state.user_type,
                    conversation_state.contract_type,
                    selected_terms=[]
                )

                conversation_state.reset()
                response = "임대차계약서에 추가하면 좋을 특약 사항의 유형들을 나열해 봤어요. 추가하고 싶은 유형을 모두 선택해 주세요!"
                options = [
                    "전세 사기 예방", "수리와 유지보수", "시설물 사용", "임대료 연체",
                    "중도 해지", "시설 변경", "애완동물 및 흡연", "보증금 반환 조건",
                    "화재 및 안전", "관리비 및 공과금", "계약 갱신"
                ]
                return jsonify({"response": response, "options": options, "next_step": "special_terms"})

            except Exception as e:
                logger.error(f"Contract generation failed: {str(e)}")
                logger.error(traceback.format_exc())
                return jsonify({"error": "계약서 생성 중 오류가 발생했습니다. 다시 시도해주세요."})

    except Exception as e:
        logger.error(f"Unexpected error: {str(e)}")
        logger.error(traceback.format_exc())
        conversation_state.reset()
        return jsonify({"error": "처리 중 오류가 발생했습니다. 다시 시도해주세요."})

    return jsonify({"response": response, "options": options})



@app.route("/download/<filename>")
def download_file(filename):
    try:
        file_path = os.path.join("generated", filename)
        if not os.path.exists(file_path):
            abort(404)
        return send_file(file_path, as_attachment=True)
    except Exception as e:
        logger.error(f"Download failed: {str(e)}")
        abort(500)

if __name__ == "__main__":
    app.run(debug=True)